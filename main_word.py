import datetime
from posixpath import splitext
import cv2 as cv
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from shutil import copy
from docx import Document
from docx.shared import Inches

import progressbar 


from utilities.exif.exif import ExifTool
from utilities.clone_detect.clone import CloneDetector
from utilities.error_level_analysis.ela import ela
from utilities.noise_analysis.plane_level import PlanesAnalysis
from utilities.noise_analysis.min_max import minmax_noise_analysis
from utilities.noise_analysis.signal_separation import analyze_signal_noise
from utilities.noise_analysis.median_noise import median_noise_inconsistencies

import os

from os.path import basename

def initiaste_paths():
    current_path = os.environ['PATH']

    #adding exiftool.exe to path
    os.environ['PATH'] = 'utilities/exif/' + os.pathsep + current_path


class ImageForgeryDetection:
    
    def __init__(self, image_path, image_name, current_time):
        self.image_path = image_path
        self.sec_list = []
        
        self.bar = progressbar.ProgressBar(maxval=20,
                                  widgets=[progressbar.Bar('=', '[', ']'), ' ', progressbar.Percentage()])
        # self.bar.start()
        self.doc = Document()
        main_heading = f"Report of {image_name} ({current_time})"
        self.doc.add_heading(main_heading, level=1)
        self.doc.add_picture(self.image_path, width=Inches(6))
        pass
    
    def get_exif(self):
        exif = ExifTool(self.image_path)
        data = ""
        for tag, value in exif.metadata.items():
                data += f"{tag}: {value}\n"
        self.bar.update(15)
        dict = {
             "head":"EXIF Meta Data",
             "image_path":None,
             "data":data
        }
        # self.sec_list.append(dict)
        self.document_add_info(dict)
        
    def get_clones(self):
        clone = CloneDetector(self.image_path)
        path = clone.detect_cloning()
        self.bar.update(15)
        dict = {
             "head":"Clone Check",
             "image_path":path,
             "data":None
        }
        self.document_add_info(dict)


    def get_bit_wise_planes_noise(self):
        tool = PlanesAnalysis(self.image_path)
        tool.preprocess()
        planes_path = tool.process()
        self.bar.update(10)
        dict = {
             "head":"Bit Wise Plane level Noise Analysis",
             "image_path":planes_path,
             "data":None
        }
        self.document_add_info(dict)
    

    def get_median_noise(self):
        path = median_noise_inconsistencies(self.image_path)
        
        self.bar.update(10)
        dict = {
             "head":"Median noise inconsistencies",
             "image_path":path,
             "data":None
        }
        self.document_add_info(dict)


    def get_min_max_noise(self):
        path = minmax_noise_analysis(self.image_path)
        self.bar.update(10)
        dict = {
             "head":"Min Max Noise Analysis",
             "image_path":path,
             "data":None
        }
        self.document_add_info(dict)

    def get_ela(self):
        path = ela(self.image_path)
        self.bar.update(15)
        dict = {
             "head":"Error Level Analysis",
             "image_path":path,
             "data":None
        }
        self.document_add_info(dict)
    
    def document_add_info(self, data):
        self.doc.add_page_break()
        self.doc.add_heading(data["head"], level = 2)
        if data["image_path"] != None:
             self.doc.add_picture(data["image_path"], width=Inches(6))
        if data["data"] != None:
            self.doc.add_paragraph(data["data"], style='List Number')
        self.bar.update(5)

    def export_document(self):
         self.doc.save(self.image_path.split('.')[0]+'_report.docx')
         self.bar.finish()
         



if __name__ == '__main__':
        
    initiaste_paths()
    img_path = askopenfilename()
    
    file_name_ext = os.path.basename(img_path).split('/')[-1]
    file_name = file_name_ext.split('.')[0]


    #move image to current export folder
    current_time = datetime.datetime.now().strftime("%Y-%m-%d_%H:%M:%S")
    
    move_path  = os.path.abspath("exports")
    move_path += f"/{file_name}_{current_time.split('_')[0]}"
    isExist = os.path.exists(move_path)
    if not isExist:
        os.makedirs(move_path)
    move_path += "/"+file_name_ext
    print(move_path)
    copy(img_path, move_path)


    tool = ImageForgeryDetection(move_path, file_name, current_time)

    tool.get_exif()
    tool.get_clones()
    tool.get_ela()

    #note issue when bitwise plane after median and before min max
    tool.get_bit_wise_planes_noise()
    tool.get_median_noise()    
    tool.get_min_max_noise()
    tool.export_document()



